from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


def _set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)


def _add_code_paragraph(doc: Document, line: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(line.rstrip("\n"))
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def md_to_docx(md_text: str, out_path: Path) -> None:
    doc = Document()
    _set_doc_defaults(doc)

    in_code = False

    # Normalize newlines
    lines = md_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")

    for raw in lines:
        line = raw.rstrip()

        # Code fence
        if line.strip().startswith("```"):
            in_code = not in_code
            if in_code:
                doc.add_paragraph("Código:")
            continue

        if in_code:
            _add_code_paragraph(doc, line)
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # Word supports heading levels 0..9; map # -> Heading 1, etc.
            doc.add_heading(text, level=min(level, 4))
            continue

        # Horizontal rule
        if re.match(r"^---+\s*$", line):
            doc.add_paragraph(" ")
            continue

        # Checkbox / bullet
        m = re.match(r"^\s*-\s*\[( |x|X)\]\s+(.*)$", line)
        if m:
            checked = m.group(1).strip().lower() == "x"
            text = m.group(2).strip()
            prefix = "☑" if checked else "☐"
            doc.add_paragraph(f"{prefix} {text}", style="List Bullet")
            continue

        # Bullet
        m = re.match(r"^\s*-\s+(.*)$", line)
        if m:
            text = m.group(1).strip()
            doc.add_paragraph(text, style="List Bullet")
            continue

        # Numbered list
        m = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m:
            text = m.group(1).strip()
            doc.add_paragraph(text, style="List Number")
            continue

        # Blank
        if not line.strip():
            doc.add_paragraph("")
            continue

        # Plain paragraph
        doc.add_paragraph(line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    md_path = root / "MANUAL_TESTES_CLIENTE.md"
    out_path = root / "MANUAL_TESTES_CLIENTE.docx"

    md_text = md_path.read_text(encoding="utf-8")
    md_to_docx(md_text, out_path)
    print(f"OK: gerado {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
